import docx

def getTextWord(wordFileName):
    doc = docx.Document(wordFileName)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


if __name__ == '__main__':

    print(getTextWord('my_doc.docx'))
    print("************************")
    doc = docx.Document('my_doc.docx')
    print("Number of paragraphs: ", len(doc.paragraphs))
    print("Paragraph 2: ", doc.paragraphs[1].text)
    print("Paragraph 2 style: ", doc.paragraphs[1].style)
    print("******runs**************")
    print("Paragraph 1: ", doc.paragraphs[0].text)
    print("Number of runs in paragraph 1: ", len(doc.paragraphs[0].runs))

    for idx, run in enumerate(doc.paragraphs[0].runs):
        print("Run %s: %s" %(idx, run.text))

    print("is Run 0 underlined: ", doc.paragraphs[0].runs[1].underline)
    print("is Run 0 underlined: ", doc.paragraphs[0].runs[1].bold)
    print("is Run 0 underlined: ", doc.paragraphs[0].runs[3].italic)
